from tkinter import *            
from tkinter import messagebox
import datetime
import sqlite3
import docx

# Create an instance of tkinter frame or ap
ap=Tk()
ap.title("Student Management Screen")
ap.iconbitmap('img/students.ico')
ap.geometry('1250x700+5+2')
ap.configure(bg='white')
photo2=PhotoImage(file='img/dashh.png')
background = Label(ap, image = photo2).place(x=0, y=0) #since tkinter doesn't support background image, we place it as a label


#background image of app
dash=PhotoImage(file='img/dashboard3.png') 
Label(ap,image=dash,bg='white').place(x=210,y=0) 

# def security():


#Function creation for dashboard button
def home():
    ho=messagebox.askyesno('Info','Do want to insert student data?')
    if (ho):
        ap.destroy()
        import newsys
image5=PhotoImage(file='img/ho.png') #addressing img
Label(ap,image=image5,bg='#5e17eb').place(x=10,y=200) #placing img
Label.image= image5 #Image was not showing so internal image varialble to object
Button(ap,width=10,pady=7,text='Dashboard',bg='#5e17eb',fg='white',font=('roboto',12),border=0,command=home).place(x=50,y=205) 



def logout():
    l=messagebox.askyesno('logout','Do you really want to logout ?')
    if (l):
        ap.destroy()
        import login
image_ext=PhotoImage(file='img/logout.png') #addressing img
Label(ap,image=image_ext,bg='#5e17eb').place(x=10,y=350) #placing img
Label.image= image_ext #Image was not showing so internal image varialble to object
Button(ap,width=5,pady=7,text='Logout',bg='#5e17eb',fg='white',font=('roboto',12),border=0,command=logout).place(x=49,y=352) 

#Function creation for exit button
def exits():
    q=messagebox.askyesno('Exit','Do you really want to exit ?')
    if (q):
        ap.destroy()
image7=PhotoImage(file='img/exit1.png') #addressing img
Label(ap,image=image7,bg='#5e17eb').place(x=10,y=550) #placing img
Label.image= image7 #Image was not showing so internal image varialble to object
Button(ap,width=5,pady=7,text='Exit',bg='#5e17eb',fg='white',font=('roboto',12),border=0,command=exits).place(x=49,y=552) 


#Function creation for student information button


#function creation for account



# Frame(ap,height=40,width=700,bg='red').place(x=205,y=170)

std=PhotoImage(file='img/stdinf.png') 
Label(ap,image=std,bg='white').place(x=202,y=170) 

def cross():
    con=sqlite3.connect('students_data.db')
    b=cid.get()
    if (cid.get()==""):
        messagebox.showerror("error","please enter student id")
    else:
        c=con.cursor()
        c.execute("DELETE from students WHERE oid=:cid",{'cid':b})
        messagebox.showinfo('success',"deleted successfully")
    con.commit()
    con.close()
    cid.delete(0,END)
    reset()
try:
    #creating a student table
    conn=sqlite3.connect('students_data.db')
    c=conn.cursor()
    c.execute("""CREATE TABLE students(
        firstname text,
        lastname text,
        gender text,
        dob int,
        mob int,
        email text,
        address text,
        course text,
        hobby text,
        ffullname text,
        foccupation text,
        feducation text,
        fmobile int,
        mfullname text,
        moccupation text,
        meducation text,
        mmobile int
    )""" )
    conn.commit()
    conn.close()
except:
    pass

#reset function
def reset():
    firstname.delete(0,END)
    lastname.delete(0,END)
    gender.delete(0,END)
    dob.delete(0,END)
    mob.delete(0,END)
    email.delete(0,END)
    address.delete(0,END)
    course.delete(0,END)
    hobby.delete(0,END)
    ffullname.delete(0,END)
    foccupation.delete(0,END)
    feducation.delete(0,END)
    fmobile.delete(0,END)
    mfullname.delete(0,END)
    moccupation.delete(0,END)
    meducation.delete(0,END)
    mmobile.delete(0,END)


# #fetch data function
def fetch():
    reset()
    a=cid.get()
    if a=="":
        messagebox.showerror("Fetch","Enter studentID")
    else:
        try:
            #database connection
            conn=sqlite3.connect('students_data.db')
            c=conn.cursor()
            c.execute("SELECT * from students where oid=:cid",{'cid':a})
            rec=c.fetchall()
            #inserting values into entry boxes
            firstname.insert(0,rec[0][0])
            lastname.insert(0,rec[0][1])
            gender.insert(0,rec[0][2])
            dob.insert(0,rec[0][3])
            mob.insert(0,rec[0][4])
            email.insert(0,rec[0][5])
            address.insert(0,rec[0][6])
            course.insert(0,rec[0][7])
            hobby.insert(0,rec[0][8])
            ffullname.insert(0,rec[0][9])
            foccupation.insert(0,rec[0][10])
            feducation.insert(0,rec[0][11])
            fmobile.insert(0,rec[0][12])
            mfullname.insert(0,rec[0][13])
            moccupation.insert(0,rec[0][14])
            meducation.insert(0,rec[0][15])
            mmobile.insert(0,rec[0][16])
            conn.commit()
            conn.close()
            #update button status to normal
            upd.config(state=NORMAL)
        except:
            messagebox.showerror("Fetch","Invalid studentID")

#submit function
def submit():
    #address values to database
    conn=sqlite3.connect('students_data.db')
    c=conn.cursor()
    c.execute("INSERT INTO students VALUES (:firstname, :lastname, :gender, :dob, :mob, :email, :address, :course, :hobby, :ffullname, :foccupation, :feducation, :fmobile, :mfullname, :moccupation, :meducation, :mmobile)",
        {
            'firstname':firstname.get(),
            'lastname':lastname.get(),
            'gender':gender.get(),
            'dob':dob.get(),
            'mob':mob.get(),
            'email':email.get(),
            'address':address.get(),
            'course':course.get(),
            'hobby':hobby.get(),
            'ffullname':ffullname.get(),
            'foccupation':foccupation.get(),
            'feducation':feducation.get(),
            'fmobile':fmobile.get(),
            'mfullname':mfullname.get(),
            'moccupation':moccupation.get(),
            'meducation':meducation.get(),
            'mmobile':mmobile.get()
        })
    conn.commit()
    #display student id
    messagebox.showinfo("info","data inserted Successfully, studentID")
    conn.commit()
    conn.close()

    #update table

    #reset entries
    reset()

  

#verification for student update
def verifyforupdate():
    #getting all occupied rooms and addressing to a list
    conn=sqlite3.connect('students_data.db')
    c=conn.cursor()
    # c.execute("SELECT Room_Number from room WHERE Room_Status=:oc",{'oc':"Occupied"})
    # list1=c.fetchall()
    # y=[]
    # for i in list1:
    #     y.append(i[0])
    conn.commit()
    conn.close()

    #getting values to verify
    a=firstname.get()
    b=lastname.get()
    c=gender.get()
    d=dob.get()
    e=mob.get()
    f=email.get()
    g=address.get()
    h=course.get()
    j=hobby.get()
    k=ffullname.get()
    l=foccupation.get()
    m=feducation.get()
    n=fmobile.get()
    o=mfullname.get()
    p=moccupation.get()
    q=meducation.get()
    r=mmobile.get()
    #verification
    if a=="" or b=="" or c=="" or d=="" or e=="" or f=="" or g=="" or h=="" or j=="" or k=="" or l=="" or m=="" or n=="" or o=="" or p=="" or q=="" or r=="":
        messagebox.showerror("error","One or More Fields Empty!")
    elif len(d)!=4:
        messagebox.showerror("error","Invalid Date")
    elif len(e or n or r)!=10:
        messagebox.showerror("error","Invalid Phone Number")
    elif "@" and ".com" not in f:
        messagebox.showerror("error","Invalid Email")
    # elif j!="T1" and j!="T2" and j!="T3" and j!="C1" and j!="C2" and j!="R1" and j!="R2" and j!="R3" and j!="R4":
    #     messagebox.showerror("error","Invalid Room Number")
    elif d[0].isalpha() or d[1].isalpha() or d[2].isalpha() or d[3].isalpha():
        messagebox.showerror("error","Invalid Date")
    elif e[0].isalpha() or e[1].isalpha() or e[2].isalpha() or e[3].isalpha() or e[4].isalpha() or e[5].isalpha() or e[6].isalpha() or e[7].isalpha() or e[8].isalpha() or e[9].isalpha():
        messagebox.showerror("error","Invalid Phone Number")
    elif n[0].isalpha() or n[1].isalpha() or n[2].isalpha() or n[3].isalpha() or n[4].isalpha() or n[5].isalpha() or n[6].isalpha() or n[7].isalpha() or n[8].isalpha() or n[9].isalpha():
        messagebox.shownrror("nrror","Invalid Phone Number")
    elif r[0].isalpha() or r[1].isalpha() or r[2].isalpha() or r[3].isalpha() or r[4].isalpha() or r[5].isalpha() or r[6].isalpha() or r[7].isalpha() or r[8].isalpha() or r[9].isalpha():
        messagebox.showerror("rrror","invalid Phone number")
    else:
        update()

#update function           
def update():
    conn=sqlite3.connect('students_data.db')
    c=conn.cursor()
    c.execute("""UPDATE students SET
        firstname=:a,
        lastname=:b,
        gender=:c,
        dob=:d,
        mob=:e,
        email=:f,
        address=:g,
        course=:h,
        hobby=:j,
        ffullname=:k,
        foccupation=:l,
        feducation=:m,
        fmobile=:n,
        mfullname=:o,
        moccupation=:p,
        meducation=:q,
        mmobile=:r
        WHERE oid=:cid""",{
            'a':firstname.get(),
            'b':lastname.get(),
            'c':gender.get(),
            'd':dob.get(),
            'e':mob.get(),
            'f':email.get(),
            'g':address.get(),
            'h':course.get(),
            'j':hobby.get(),
            'k':ffullname.get(),
            'l':foccupation.get(),
            'm':feducation.get(),
            'n':fmobile.get(),
            'o':mfullname.get(),
            'p':moccupation.get(),
            'q':meducation.get(),
            'r':mmobile.get(),
            'cid':cid.get()
        })
    conn.commit()
    conn.close()
    reset()

    messagebox.showinfo("Update","Data Updated Successfully")

#verification for submitting
def verifyforsubmit():
    conn=sqlite3.connect('students_data.db')
    c=conn.cursor()
    conn.commit()
    conn.close()
    a=firstname.get()
    b=lastname.get()
    c=gender.get()
    d=dob.get()
    e=mob.get()
    f=email.get()
    g=address.get()
    h=course.get()
    j=hobby.get()
    k=ffullname.get()
    l=foccupation.get()
    m=feducation.get()
    n=fmobile.get()
    o=mfullname.get()
    p=moccupation.get()
    q=meducation.get()
    r=mmobile.get()
    if a=="" or b=="" or c=="" or d=="" or e=="" or f=="" or g=="" or h=="" or j=="" or k=="" or l=="" or m=="" or n=="" or o=="" or p=="" or q=="" or r=="":
        messagebox.showerror("Error","One or More Fields Empty!")
    elif len(d)!=4:
        messagebox.showerror("Error","Invalid Date")
    elif len(e or n or r)!=10:
        messagebox.showerror("Error","Invalid Phone Number")
    elif "@" and ".com" not in f:
        messagebox.showerror("Error","Invalid Email")
    elif d[0].isalpha() or d[1].isalpha() or d[2].isalpha() or d[3].isalpha():
        messagebox.showerror("Error","Invalid Date")
    elif e[0].isalpha() or e[1].isalpha() or e[2].isalpha() or e[3].isalpha() or e[4].isalpha() or e[5].isalpha() or e[6].isalpha() or e[7].isalpha() or e[8].isalpha() or e[9].isalpha():
        messagebox.showerror("Error","Invalid Phone Number")
    elif n[0].isalpha() or n[1].isalpha() or n[2].isalpha() or n[3].isalpha() or n[4].isalpha() or n[5].isalpha() or n[6].isalpha() or n[7].isalpha() or n[8].isalpha() or n[9].isalpha():
        messagebox.shownrror("nrror","Invalid Phone Number")
    elif r[0].isalpha() or r[1].isalpha() or r[2].isalpha() or r[3].isalpha() or r[4].isalpha() or r[5].isalpha() or r[6].isalpha() or r[7].isalpha() or r[8].isalpha() or r[9].isalpha():
        messagebox.showerror("rrror","invalid Phone number")
    else:
        submit()

#Labels for data entry


# Label(ap,text='Insert student data',bg='#DCEF01',fg='grey',font=('roboto',16,'bold')).place(x=600,y=100)
Label(ap,text="Student id:",fg='black', bg='white',font=('roboto',10,'bold')).place(x=210,y=120)
Label(ap,text="First Name",fg='black', bg='white',font=('roboto',10,'bold')).place(x=210,y=220)
Label(ap,text="Last Name",fg='black', bg='white',font=('roboto',10,'bold')).place(x=390,y=220)
Label(ap,text="Gender",fg='black', bg='white',font=('roboto',10,'bold')).place(x=565,y=220)
Label(ap,text="Year of Birth",fg='black', bg='white',font=('roboto',10,'bold')).place(x=740,y=220)
Label(ap,text="Mobile",fg='black', bg='white',font=('roboto',10,'bold')).place(x=210,y=277)
Label(ap,text="Email",fg='black', bg='white',font=('roboto',10,'bold')).place(x=390,y=277)
Label(ap,text="Address",fg='black', bg='white',font=('roboto',10,'bold')).place(x=565,y=277)
Label(ap,text="Course",fg='black', bg='white',font=('roboto',10,'bold')).place(x=740,y=277)
Label(ap,text="Hobby",fg='black', bg='white',font=('roboto',10,'bold')).place(x=210,y=337)
#Entry boxes
cid=Entry(ap,width=30,fg='black',border=0,highlightthickness=2,highlightbackground='#F1F0FD',font=('roboto',9))
firstname=Entry(ap,width=15,fg='black',border=0,highlightthickness=2,highlightbackground='#F1F0FD',font=('roboto',9))
lastname=Entry(ap,width=15,fg='black',border=0,highlightthickness=2,highlightbackground='#F1F0FD',font=('roboto',9))
gender=Entry(ap,width=15,fg='black',border=0,highlightthickness=1,highlightbackground='#F1F0FD',font=('roboto',9))
dob=Entry(ap,width=15,fg='black',border=0,highlightthickness=1,highlightbackground='#F1F0FD',font=('roboto',9))
mob=Entry(ap,width=15,fg='black',border=0,highlightthickness=1,highlightbackground='#F1F0FD',font=('roboto',9))
email=Entry(ap,width=15,fg='black',border=0,highlightthickness=1,highlightbackground='#F1F0FD',font=('roboto',9))
address=Entry(ap,width=15,fg='black',border=0,highlightthickness=1,highlightbackground='#F1F0FD',font=('roboto',9))
course=Entry(ap,width=15,fg='black',border=0,highlightthickness=1,highlightbackground='#F1F0FD',font=('roboto',9))
hobby=Entry(ap,width=15,fg='black',border=0,highlightthickness=1,highlightbackground='#F1F0FD',font=('roboto',9))
cid.place(x=290,y=120,height=25,width=155)
firstname.place(x=210,y=240,height=25,width=155)
lastname.place(x=390,y=240,height=25,width=155)
gender.place(x=565,y=240,height=25,width=155)
dob.place(x=740,y=240,height=25,width=155)
mob.place(x=210,y=297,height=25,width=155)
email.place(x=390,y=297,height=25,width=155)
address.place(x=565,y=297,height=25,width=155)
course.place(x=740,y=297,height=25,width=155)
hobby.place(x=210,y=360,height=25,width=155)

#buttons
Button(ap,width=10,pady=2,text='FETCH DATA',bg='#5e17eb',fg='white',border=0,cursor='hand2',command=fetch).place(x=460,y=120)


#father label and box
Frame(ap,height=600,width=700,bg='white').place(x=210,y=390)
fathers=PhotoImage(file='img/father.png') 
Label(image=fathers,bg='white').place(x=203,y=390) 
Label(text="full Name",fg='black', bg='white',font=('roboto',10,'bold')).place(x=210,y=440)
ffullname=Entry(width=15,fg='black',border=0,highlightthickness=2,highlightbackground='#F1F0FD',font=('roboto',9))
ffullname.place(x=210,y=465,height=25,width=155)

Label(text="Occupation",fg='black', bg='white',font=('roboto',10,'bold')).place(x=390,y=440)
foccupation=Entry(width=15,fg='black',border=0,highlightthickness=2,highlightbackground='#F1F0FD',font=('roboto',9))
foccupation.place(x=390,y=465,height=25,width=155)


Label(text="Education",fg='black', bg='white',font=('roboto',10,'bold')).place(x=560,y=440)
feducation=Entry(width=15,fg='black',border=0,highlightthickness=2,highlightbackground='#F1F0FD',font=('roboto',9))
feducation.place(x=565,y=465,height=25,width=155)


Label(text="Mobile No.",fg='black', bg='white',font=('roboto',10,'bold')).place(x=740,y=440)
fmobile=Entry(width=15,fg='black',border=0,highlightthickness=2,highlightbackground='#F1F0FD',font=('roboto',9))
fmobile.place(x=740,y=465,height=25,width=155)




#father label and box

mothers=PhotoImage(file='img/mother.png') 
Label(image=mothers,bg='white').place(x=203,y=500) 
Label(text="full Name",fg='black', bg='white',font=('roboto',10,'bold')).place(x=210,y=550)
mfullname=Entry(width=15,fg='black',border=0,highlightthickness=2,highlightbackground='#F1F0FD',font=('roboto',9))
mfullname.place(x=210,y=573,height=25,width=155)

Label(text="Occupation",fg='black', bg='white',font=('roboto',10,'bold')).place(x=390,y=550)
moccupation=Entry(width=15,fg='black',border=0,highlightthickness=2,highlightbackground='#F1F0FD',font=('roboto',9))
moccupation.place(x=390,y=573,height=25,width=155)


Label(text="Education",fg='black', bg='white',font=('roboto',10,'bold')).place(x=560,y=550)
meducation=Entry(width=15,fg='black',border=0,highlightthickness=2,highlightbackground='#F1F0FD',font=('roboto',9))
meducation.place(x=565,y=573,height=25,width=155)


Label(text="Mobile No.",fg='black', bg='white',font=('roboto',10,'bold')).place(x=740,y=550)
mmobile=Entry(width=15,fg='black',border=0,highlightthickness=2,highlightbackground='#F1F0FD',font=('roboto',9))
mmobile.place(x=740,y=573,height=25,width=155)

Button(ap,width=10,pady=7,text='SUBMIT',bg='#5e17eb',fg='white',border=0,cursor='hand2',command=verifyforsubmit).place(x=1000,y=430)
upd=Button(ap,width=10,pady=7,text='UPDATE',bg='#5e17eb',fg='white',border=0,command=verifyforupdate,state=DISABLED,cursor='hand2')
upd.place(x=1000, y=500)
Button(ap,width=10,pady=7,text='CLEAR',bg='#5e17eb',fg='white',border=0,command=reset,cursor='hand2').place(x=1100, y=430)
Button(ap,width=10,pady=7,text='Delete',bg='#5e17eb',fg='white',border=0,command=cross,cursor='hand2').place(x=1100, y=500)


#calling table function


#--------------result------------------------------#
def result():
    def calculate_result():
        if entry1.get()=="" or entry2.get()=="" or entry3.get()=="" or entry4.get()=="":
            messagebox.showerror('error','fields is missing')
        else:
        # Get values from input fields
            subject1 = float(entry1.get())
            subject2 = float(entry2.get())
            subject3 = float(entry3.get())

        # Calculate total marks and percentage
        total_marks = subject1 + subject2 + subject3
        percentage = (total_marks / 300) * 100

        # Determine grade based on percentage
        if percentage >= 90:
            grade = "A+"
        elif percentage >= 80:
            grade = "A"
        elif percentage >= 70:
            grade = "B"
        elif percentage >= 60:
            grade = "C"
        elif percentage >= 50:
            grade = "D"
        else:
            grade = "F"

        # Update result labels
        total_marks_label.config(text="Total marks obtained: {:.2f}".format(total_marks))
        percentage_label.config(text="Percentage obtained: {:.2f}%".format(percentage))
        grade_label.config(text="Grade: {}".format(grade))
        entry4.config(text='Name of student: {}'.format(entry4.get()))
        entry1.config(text='Python: {}'.format(entry1.get()))
        entry2.config(text='math: {}'.format(entry2.get()))
        entry3.config(text='software design: {}'.format(entry3.get()))

    def print_result():
        if entry1.get()=="" or entry2.get()=="" or entry3.get()=="" or entry4.get()=="":
            messagebox.showerror('error','fields is missing')
        else:
        # Get values from result labels
            total_marks = total_marks_label.cget("text")
            percentage = percentage_label.cget("text")
            grade = grade_label.cget("text")
            py=entry1.cget("text")
            mt=entry2.cget("text")
            sofd=entry3.cget("text")
            name=entry4.cget("text")

            # Create a new document
            document = docx.Document()

            # Add a heading
            document.add_heading("-----------------------------------------Result------------------------------------------------")
            document.add_paragraph('remark:')
            document.add_paragraph(name)
            document.add_paragraph(py)
            document.add_paragraph(mt)
            document.add_paragraph(sofd)
            # Add the result information
            document.add_paragraph(total_marks)
            document.add_paragraph(percentage)
            document.add_paragraph(grade)

            # Save the document with timestamp in the file name
            timestamp = datetime.datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
            document.save("result_{}, {}.docx".format(timestamp,entry4.get()))
            messagebox.showinfo("Success","Result is created as docx")

        # Create the main ap


        # Create input fields and labels


    frame_1=Frame(ap,height=900,width=1100,bg='white').place(x=200,y=0)

    header=Label(frame_1,text='Result Generator',fg='#3751FE', bg='white',font=('Gill Sans MT',23,'bold'))
    header.place(x=600,y=10)
    # header=Label(frame_1,text='Enter Your Marks',fg='#3751FE', bg='white',font=('Gill Sans MT',23,'bold'))
    # header.place(x=200,y=40)
    label1 = Label(frame_1, text="Enter mark of python",fg='black', bg='white',font=('roboto',10))
    label1.place(x=250,y=160)

    entry1 = Entry(frame_1,width=20,fg='black',border=0,highlightthickness=2,highlightbackground='#F1F0FD',font=('roboto',9))
    entry1.place(x=250,y=180,height=30)

    label2 = Label(frame_1, text="Enter mark of math",fg='black', bg='white',font=('roboto',10))
    label2.place(x=250,y=210)

    entry2 = Entry(frame_1,width=20,fg='black',border=0,highlightthickness=2,highlightbackground='#F1F0FD',font=('roboto',9))
    entry2.place(x=250,y=230,height=30)

    label3 = Label(frame_1, text="Enter marks of software design",fg='black', bg='white',font=('roboto',10))
    label3.place(x=250,y=270)

    entry3 = Entry(frame_1,width=20,fg='black',border=0,highlightthickness=2,highlightbackground='#F1F0FD',font=('roboto',9))
    entry3.place(x=250,y=290,height=30)

    label4 = Label(frame_1, text="Enter name of student",fg='black', bg='white',font=('roboto',10))
    label4.place(x=250,y=100)

    entry4 = Entry(frame_1,width=20,fg='black',border=0,highlightthickness=2,highlightbackground='#F1F0FD',font=('roboto',9))
    entry4.place(x=250,y=120,height=30)

    # Create calculate button
    calculate_button = Button(frame_1, text="Calculate",width=10,pady=7,bg='#3751FE',fg='white',border=0,command=calculate_result)
    calculate_button.place(x=250,y=340)

    # Create print button
    print_button = Button(frame_1, text="Print",width=10,pady=7,bg='#3751FE',fg='white',border=0,command=print_result)
    print_button.place(x=350,y=340)

    # Create result labels
    total_marks_label =Label(frame_1, text="Total Marks obtained",fg='black', bg='white',font=('roboto',10))
    total_marks_label.place(x=250,y=390)

    percentage_label = Label(frame_1, text="Percentage obtained",fg='black', bg='white',font=('roboto',10))
    percentage_label.place(x=250,y=420)

    grade_label = Label(frame_1, text="Grade obtained",fg='black', bg='white',font=('roboto',10))
    grade_label.place(x=250,y=450)

    image24=PhotoImage(file='img/resultss.png') #addressing img
    Label(frame_1,image=image24,bg='white').place(x=600,y=100) #placing img
    Label.image= image24
image13=PhotoImage(file='img/result.png') #addressing img
Label(ap,image=image13,bg='#5e17eb').place(x=10,y=300) #placing img
Label.image= image13 #Image was not showing so internal image varialble to object
Button(ap,width=5,pady=7,text='Result',bg='#5e17eb',fg='white',font=('roboto',12),border=0,command=result).place(x=49,y=302) 








def table():
    
#value of y for search function
    space=203

    #search function
    def srch():
        #get data from dropdown menu and entry
        term=clicked.get()
        detail=search.get()

        #database connection
        conn=sqlite3.connect('students_data.db')
        c=conn.cursor()
        c.execute("SELECT oid,* from students")
        det=c.fetchall()

        #Searching algorithm
        i=len(det)-1
        while i>=0:
            #searching for Student id
            if term=="Student ID":
                try:
                    #comparing values
                    int(detail)
                    if det[i][0]!=int(detail):
                        i=i-1
                        if i==-1:
                            break
                    else:
                        Label(ap,text="-->",bg='white',font=('Agency FB',13,'bold'),fg='red').place(x=200,y=(space+20*i))
                        break
                except:
                    messagebox.showerror("Search","Invalid student ID")
                    break
            #searching by first name
            elif term=="First Name":
                if det[i][1]!=detail:
                    i=i-1
                    if i==-1:
                        break
                else:
                    Label(ap,text="-->",bg='white',font=('Agency FB',13,'bold'),fg='red').place(x=200,y=(space+20*i))
                    i=i-1
                    if i==-1:
                        break
            #searching by last name
            elif term=="Last Name":
                if det[i][2]!=detail:
                    i=i-1
                    if i==-1:
                        break
                else:
                    Label(ap,text="-->",bg='white',font=('Agency FB',13,'bold'),fg='red').place(x=200,y=(space+20*i))
                    i=i-1
                    if i==-1:
                        break
            #searching by mobile
            elif term=="Mobile":
                if det[i][5]!=detail:
                    i=i-1
                    if i==-1:
                        break
                else:
                    Label(ap,text="-->",bg='white',font=('Agency FB',13,'bold'),fg='red').place(x=200,y=(space+20*i))
                    break
            #searching by email
            elif term=="Email":
                if det[i][6]!=detail:
                    i=i-1
                    if i==-1:
                        break
                else:
                    Label(ap,text="-->",bg='white',font=('Agency FB',13,'bold'),fg='red').place(x=200,y=(space+20*i))
                    break
            else:
                pass
        conn.commit()
        conn.close()    

    data_2=PhotoImage(file='img/dashboard3.png') 
    Label(ap,image=data_2,bg='white').place(x=200,y=0)
    image33=PhotoImage(file='img/stw.png') #addressing img
    Label(ap,image=image33,bg='white').place(x=200,y=0) #placing img
    Label.image= image33 #Image was not showing so internal image varialble to object
    #label, dropdown menu for searching
    Label(ap,text='Search',bg='white',font=('roboto',12)).place(x=210,y=120)
    clicked=StringVar()
    clicked.set("Student ID")
    drop=OptionMenu(ap,clicked,"Student ID","First Name","Last Name","Mobile","Email")
    drop.place(x=290,y=120)
    drop.config(width=15,border=0,background='white')

    #entry for searching
    search=Entry(ap,border=0,highlightthickness=3)
    search.place(x=290,y=150,height=28,width=145)
    Button(ap, text="Search",font=('Arial',8,'bold'),fg='white',bg="purple",width=8,height=1,border=0,cursor='hand2',command=srch).place(x=440,y=120)
    # Label(ap,text='Students Records',fg='#3751FE', bg='white',font=('Gill Sans MT',23,'bold')).place(x=600,y=10)
    # para=Label(ap,text='Please donot misuse the information.',fg='#ef3a5d', bg='white',font=('roboto',12,'bold'))
    # para.place(x=580,y=60)
    #table
    def tbl():
        table=Frame(ap,height=580,width=950,bg='white')
        table.place(x=220,y=190)
       

        try:
            #try fetching data from database
            conn=sqlite3.connect('students_data.db')
            c=conn.cursor()
            c.execute("SELECT oid, firstname, lastname, gender, dob, mob, email, address, course, hobby, ffullname, fmobile, mfullname, mmobile from students")
            lst=c.fetchall()
            conn.commit()
            conn.close()
        except:
            #empty list if list doesn't exist
            lst=[]
        finally:
            #Table headings
            lst.insert(0,('ID','First Name','Last Name','Gender','Dob','Mobile','Email','Address','course','hobby','ffullname','fmobile','mfullname','mmobile'))

        #creating a table
        total_rows =len(lst)
        total_columns=len(lst[0])
        for i in range(total_rows):
            if i==0:
                #table heading
                fontt=('roboto',10,'bold')
                jus=CENTER
                bgc ='white'
            else:
                #table data
                fontt=('Arial',10)
                jus=LEFT
                bgc='white'
            for j in range(total_columns):
                #width for all columns
                if j==0:
                    wid=3
                elif j==1 or j==2:
                    wid=10
                elif j==3:
                    wid=7
                elif j==4:
                    wid=7
                elif j==5:
                    wid=10
                elif j==6:
                    wid=13
                elif j==7:
                    wid=9
                elif j==8:
                    wid=10
                elif j==9:
                    wid=9
                elif j==10:
                    wid=15
                elif j==11:
                    wid=11
                elif j==12:
                    wid=15
                elif j==13:
                    wid=11
                else:
                    wid=8
                e=Entry(
                    table,
                    width=wid,
                    font=fontt,
                    justify=jus,
                    disabledforeground='black',
                    disabledbackground=bgc
                )
                e.grid(row=i,column=j)
                e.insert(0,lst[i][j])
                e.config(state=DISABLED)

    #calling table function
    tbl()
def student():
    ab=messagebox.askyesno('Account','view student data')
    if (ab):
        ap.destroy()
        import information
image6=PhotoImage(file='img/students.png') #addressing img
Label(ap,image=image6,bg='#5e17eb').place(x=10,y=250) #placing img
Label.image= image6 #Image was not showing so internal image varialble to object
Button(ap,width=10,pady=7,text='Students',bg='#5e17eb',fg='white',font=('roboto',12),border=0,command=table).place(x=42,y=252) 


ap.mainloop()

